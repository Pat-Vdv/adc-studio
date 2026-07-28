"""Renderer DOCX — matérialise un Document (IR) en fichier Word.

Le renderer est **maître de la mise en page** : il construit le document
directement depuis l'IR avec python-docx (aucun template rempli par
substitution). Développement incrémental — composants rendus :
  - C-001-cover
  - C-002-identity-page
  - C-003-executive-summary
  - C-009-environment
  - C-008-timeline
  - C-004-finding
  - C-007-decision
  - C-005-recommendation
  - C-006-risk

Un composant présent dans l'IR mais sans renderer est **ignoré proprement**
(pas d'exception, pas de contenu fantôme) : la traçabilité des composants non
rendus reste portée par `Document.diagnostics`, jamais masquée.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Callable

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .model import ComponentInstance, Document

# Un renderer reçoit (docx, instance, contexte) et écrit dans le document Word
# en place. Le contexte est le `render_context` de l'IR (ADR-0008) : index
# techniques communs au document, par exemple les libellés des cibles
# référencées par identifiant.
Renderer = Callable[[Any, ComponentInstance, dict[str, Any]], None]


# Libellés de présentation des énumérations partagées par plusieurs composants
# (gravité d'un constat, priorité d'une recommandation). Le modèle conserve les
# valeurs canoniques anglaises : la traduction n'existe que dans cette couche.
_ENUM_LABELS: dict[str, str] = {
    "low": "Faible",
    "medium": "Moyenne",
    "high": "Élevée",
    "critical": "Critique",
}


def _enum_label(value: Any) -> Any:
    """Libellé lisible d'une valeur d'énumération partagée.

    Une valeur hors vocabulaire est restituée telle quelle : mieux vaut afficher
    la valeur canonique de la source qu'inventer une traduction.
    """
    return _ENUM_LABELS.get(value, value) if isinstance(value, str) else value


def _add_label_value(docx: Any, label: str, value: Any) -> None:
    """Ligne « Label : valeur », omise si la valeur est absente."""
    if value in (None, ""):
        return
    paragraph = docx.add_paragraph()
    run = paragraph.add_run(f"{label} : ")
    run.bold = True
    paragraph.add_run(str(value))


def _render_cover(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    document_type = payload.get("document_type")
    if document_type:
        p = docx.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(document_type).upper())
        run.bold = True
        run.font.size = Pt(14)

    title = docx.add_heading(payload.get("title") or "", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = payload.get("subtitle")
    if subtitle:
        p = docx.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(subtitle))
        run.italic = True
        run.font.size = Pt(12)

    docx.add_paragraph()  # respiration

    _add_label_value(docx, "Client", payload.get("client"))
    _add_label_value(docx, "Auteur", payload.get("author"))
    _add_label_value(docx, "Date", payload.get("date"))
    _add_label_value(docx, "Version", payload.get("version"))
    _add_label_value(docx, "Référence", payload.get("reference"))
    _add_label_value(docx, "Confidentialité", payload.get("confidentiality"))


_IDENTITY_LABELS: tuple[tuple[str, str], ...] = (
    ("id", "Identifiant"),
    ("reference", "Référence"),
    ("title", "Titre"),
    ("client", "Client"),
    ("author", "Auteur"),
    ("version", "Version"),
    ("date", "Date"),
    ("language", "Langue"),
    ("confidentiality", "Confidentialité"),
)

_IDENTITY_TABLES: tuple[tuple[str, str, tuple[tuple[str, str], ...]], ...] = (
    (
        "revisions",
        "Révisions",
        (("version", "Version"), ("date", "Date"), ("author", "Auteur"), ("summary", "Objet")),
    ),
    ("validations", "Validations", (("role", "Rôle"), ("name", "Nom"), ("date", "Date"))),
    (
        "distribution",
        "Diffusion",
        (("name", "Nom"), ("organisation", "Organisation"), ("role", "Rôle")),
    ),
)


def _add_table(docx: Any, columns: tuple[tuple[str, str], ...], rows: Any) -> None:
    """Table à en-tête, une ligne par entrée ; rien si la liste est vide."""
    if not rows:
        return
    table = docx.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, (_, header) in zip(table.rows[0].cells, columns):
        cell.text = ""
        cell.paragraphs[0].add_run(header).bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, (key, _) in zip(cells, columns):
            value = row.get(key)
            cell.text = "" if value in (None, "") else str(value)


def _render_identity_page(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    # Composant pleine page : il démarre après la page de garde.
    docx.add_page_break()
    docx.add_heading(payload.get("heading") or "Identité du document", level=1)

    identification = payload.get("identification") or {}
    for key, label in _IDENTITY_LABELS:
        _add_label_value(docx, label, identification.get(key))

    for key, heading, columns in _IDENTITY_TABLES:
        rows = payload.get(key) or ()
        if not rows:
            continue  # section absente de la source : pas de tableau vide
        docx.add_heading(heading, level=2)
        _add_table(docx, columns, rows)


_SUMMARY_SECTIONS: tuple[tuple[str, str], ...] = (
    ("context", "Contexte"),
    ("business_impact", "Impact métier"),
    ("conclusion", "Conclusion"),
    ("recommended_action", "Action recommandée"),
)


def _render_executive_summary(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    # Le résumé exécutif ouvre le corps du rapport : il démarre sur sa page.
    docx.add_page_break()
    docx.add_heading(payload.get("heading") or "Résumé exécutif", level=1)

    for key, heading in _SUMMARY_SECTIONS:
        blocks = payload.get(key) or ()
        if not blocks:
            continue  # volet non renseigné : pas de titre orphelin
        docx.add_heading(heading, level=2)
        for block in blocks:
            docx.add_paragraph(str(block))


_ENVIRONMENT_LABELS: tuple[tuple[str, str], ...] = (
    ("server_name", "Serveur"),
    ("operating_system", "Système d'exploitation"),
    ("database_engine", "Moteur de base de données"),
    ("database_engine_version", "Version du moteur"),
    ("instance", "Instance"),
    ("primary_database", "Base principale"),
    ("collation", "Collation"),
    ("cpu_logical_count", "Processeurs logiques"),
    ("memory_gb", "Mémoire (Go)"),
)

_STORAGE_COLUMNS: tuple[tuple[str, str], ...] = (
    ("volume", "Volume"),
    ("role", "Rôle"),
    ("allocation_unit_kb", "Unité d'allocation (Ko)"),
)


def _render_environment(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    # Section courante du corps : pas de saut de page, elle suit le narratif.
    docx.add_heading(payload.get("heading") or "Environnement", level=1)

    system = payload.get("system") or {}
    for key, label in _ENVIRONMENT_LABELS:
        _add_label_value(docx, label, system.get(key))

    storage = payload.get("storage") or ()
    if storage:
        docx.add_heading("Stockage", level=2)
        _add_table(docx, _STORAGE_COLUMNS, storage)


_TIMELINE_COLUMNS: tuple[tuple[str, str], ...] = (
    ("timestamp", "Horodatage"),
    ("title", "Événement"),
    ("description", "Description"),
)


def _render_timeline(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    entries = instance.payload.get("entries") or ()
    if not entries:
        return  # chronologie vide : pas de titre, pas de tableau

    # Section courante du corps, dans l'ordre de la source.
    docx.add_heading(instance.payload.get("heading") or "Chronologie", level=1)
    _add_table(docx, _TIMELINE_COLUMNS, entries)


_FINDING_SECTIONS: tuple[tuple[str, str], ...] = (
    ("observation", "Observation"),
    ("impact", "Impact"),
    ("analysis", "Analyse"),
)


def _reference_labels(references: Any, context: dict[str, Any], index: str) -> tuple[str, ...]:
    """Libellés lisibles des cibles référencées par identifiant.

    Les identifiants sont des clés de liaison internes : ils ne sont jamais
    écrits dans le document. Une référence dont le libellé est introuvable est
    donc omise — l'intégrité référentielle est garantie en amont par la
    validation de la source.
    """
    titles = context.get(index) or {}
    return tuple(titles[ref] for ref in references or () if ref in titles)


def _render_finding(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    # Chaque constat est une section autonome : le renderer reste sans état et
    # n'émet aucun titre de partie commun aux occurrences.
    title = payload.get("title")
    docx.add_heading(f"Constat — {title}" if title else "Constat", level=1)

    _add_label_value(docx, "Gravité", _enum_label(payload.get("severity")))

    for key, heading in _FINDING_SECTIONS:
        blocks = payload.get(key) or ()
        if not blocks:
            continue  # volet non renseigné : pas de titre orphelin
        docx.add_heading(heading, level=2)
        for block in blocks:
            docx.add_paragraph(str(block))

    labels = _reference_labels(payload.get("evidence_ids"), context, "evidence_titles")
    if labels:
        _add_label_value(docx, "Preuves", " ; ".join(labels))


def _render_decision(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    # Section autonome, comme le constat : renderer sans état, pas de titre de
    # partie commun aux occurrences.
    title = payload.get("title")
    docx.add_heading(f"Mesure prise — {title}" if title else "Mesure prise", level=1)

    # Valeur canonique de la source (« completed », …) : aucune traduction ici.
    _add_label_value(docx, "Statut", payload.get("status"))

    for block in payload.get("description") or ():
        docx.add_paragraph(str(block))


_RECOMMENDATION_SECTIONS: tuple[tuple[str, str], ...] = (
    ("description", "Description"),
    ("rationale", "Justification"),
)


def _render_recommendation(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    # Section autonome, dans l'ordre de la source : aucun regroupement ni tri
    # par priorité, aucun titre de partie commun aux occurrences.
    title = payload.get("title")
    docx.add_heading(f"Recommandation — {title}" if title else "Recommandation", level=1)

    _add_label_value(docx, "Priorité", _enum_label(payload.get("priority")))

    for key, heading in _RECOMMENDATION_SECTIONS:
        blocks = payload.get(key) or ()
        if not blocks:
            continue  # volet non renseigné : pas de titre orphelin
        docx.add_heading(heading, level=2)
        for block in blocks:
            docx.add_paragraph(str(block))

    labels = _reference_labels(payload.get("related_finding_ids"), context, "finding_titles")
    if labels:
        _add_label_value(docx, "Constats liés", " ; ".join(labels))


def _render_risk(docx: Any, instance: ComponentInstance, context: dict[str, Any]) -> None:
    payload = instance.payload

    # Section autonome, dans l'ordre de la source : ni tri, ni regroupement par
    # niveau, ni titre de partie commun aux occurrences.
    title = payload.get("title")
    docx.add_heading(f"Risque — {title}" if title else "Risque", level=1)

    _add_label_value(docx, "Niveau", _enum_label(payload.get("level")))

    for block in payload.get("description") or ():
        docx.add_paragraph(str(block))

    labels = _reference_labels(
        payload.get("mitigation_recommendation_ids"), context, "recommendation_titles"
    )
    if labels:
        _add_label_value(docx, "Traitement prévu", " ; ".join(labels))


_RENDERERS: dict[str, Renderer] = {
    "C-001-cover": _render_cover,
    "C-002-identity-page": _render_identity_page,
    "C-003-executive-summary": _render_executive_summary,
    "C-009-environment": _render_environment,
    "C-008-timeline": _render_timeline,
    "C-004-finding": _render_finding,
    "C-007-decision": _render_decision,
    "C-005-recommendation": _render_recommendation,
    "C-006-risk": _render_risk,
}


def render_docx(document: Document, output_path: str | Path) -> Path:
    """Rend le Document (IR) dans un fichier .docx et retourne son chemin.

    Les composants sans renderer sont ignorés sans erreur ; la génération ne
    dépend pas de leur prise en charge.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Contexte de présentation, commun à toutes les instances : il vient de
    # l'IR, jamais de la source. Le renderer ne connaît pas le JSON d'entrée.
    context = dict(document.metadata.get("render_context") or {})

    docx = DocxDocument()
    for instance in document.components:
        renderer = _RENDERERS.get(instance.component_id)
        if renderer is None:
            continue  # non rendu à ce stade — voir Document.diagnostics
        renderer(docx, instance, context)

    docx.save(str(output_path))
    return output_path
