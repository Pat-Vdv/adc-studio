from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
import shutil, re

ROOT = Path('/mnt/data/ADC-Studio-Sprint002')
for d in ['docs','brand','templates/word','examples','tools']:
    (ROOT/d).mkdir(parents=True, exist_ok=True)

BLUE = '0054A6'
DARK = '1F2937'
GRAY = '6B7280'
LIGHT = 'EEF4FA'
GREEN = '2E7D32'
ORANGE = 'F59E0B'
RED = 'C62828'
WHITE = 'FFFFFF'
BORDER = 'D1D5DB'


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = 'w:' + edge
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in kwargs[edge].items():
                element.set(qn('w:' + key), str(val))


def add_field(paragraph, instr, display=''):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = instr
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = display
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar, instrText, fldChar2, text, fldChar3])


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn('w:updateFields'))
    if update is None:
        update = OxmlElement('w:updateFields')
        settings.append(update)
    update.set(qn('w:val'), 'true')


def set_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    el = OxmlElement('w:keepNext')
    pPr.append(el)


def configure_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for idx, size in [(1,24),(2,17),(3,13),(4,11)]:
        st = styles[f'Heading {idx}']
        st.font.name = 'Aptos Display'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE if idx < 3 else DARK)
        st.paragraph_format.space_before = Pt(16 if idx < 3 else 10)
        st.paragraph_format.space_after = Pt(6)
        set_keep_with_next(st.element.getparent() if False else doc.add_paragraph())
        # remove accidental paragraph immediately
        doc._body._body.remove(doc.paragraphs[-1]._p)

    custom = {
        'ADC Subtitle': (16, GRAY, False),
        'ADC Lead': (12, DARK, False),
        'ADC Metadata': (8.5, GRAY, False),
        'ADC Caption': (8.5, GRAY, True),
        'ADC Code': (9, DARK, False),
        'ADC Callout Info': (10, DARK, False),
        'ADC Callout Success': (10, DARK, False),
        'ADC Callout Warning': (10, DARK, False),
        'ADC Callout Critical': (10, DARK, False),
        'ADC Quote': (11, GRAY, False),
    }
    for name,(size,color,bold) in custom.items():
        if name not in styles:
            st=styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
        else: st=styles[name]
        st.font.name = 'Consolas' if name=='ADC Code' else 'Aptos'
        st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=bold
        st.paragraph_format.space_after=Pt(5)
        if name=='ADC Code':
            st.paragraph_format.left_indent=Cm(.5); st.paragraph_format.right_indent=Cm(.5)
        if name=='ADC Quote':
            st.paragraph_format.left_indent=Cm(.8); st.font.italic=True

    # Table styles are applied manually for portability.


def configure_page(section):
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(.7)
    section.footer_distance = Cm(.7)


def add_header_footer(section):
    header=section.header
    p=header.paragraphs[0]
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run('A.D.C.  |  ADC Studio')
    r.font.name='Aptos'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    footer=section.footer
    table=footer.add_table(1,2,Cm(16.8))
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width=Cm(12.5); table.columns[1].width=Cm(4.3)
    c1,c2=table.rows[0].cells
    p1=c1.paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p1.add_run('Confidentiel - A.D.C. srl'); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    p2=c2.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p2.add_run('Page '); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    add_field(p2,'PAGE','1'); p2.add_run(' / '); add_field(p2,'NUMPAGES','1')
    # remove default empty paragraph
    if len(footer.paragraphs)>0 and not footer.paragraphs[0].text:
        footer._element.remove(footer.paragraphs[0]._p)


def add_cover(doc, title='Titre du rapport', subtitle='Sous-titre ou périmètre', client='Nom du client', version='1.0'):
    sec=doc.sections[0]
    sec.different_first_page_header_footer=True
    for _ in range(4): doc.add_paragraph('')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run('A.D.C.'); r.bold=True; r.font.name='Aptos Display'; r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(24)
    r=p.add_run(title); r.bold=True; r.font.name='Aptos Display'; r.font.size=Pt(32); r.font.color.rgb=RGBColor.from_string(DARK)
    p=doc.add_paragraph(subtitle, style='ADC Subtitle')
    p.paragraph_format.space_after=Pt(42)
    table=doc.add_table(4,2)
    table.alignment=WD_TABLE_ALIGNMENT.LEFT
    labels=['Client','Document','Version','Statut']
    vals=[client,'ADC-RPT-XXXX',version,'Brouillon']
    for i,(lab,val) in enumerate(zip(labels,vals)):
        table.cell(i,0).text=lab; table.cell(i,1).text=val
        set_cell_shading(table.cell(i,0),LIGHT)
        for c in table.rows[i].cells:
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(c,bottom={'val':'single','sz':'4','color':BORDER})
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for rr in p.runs: rr.font.size=Pt(9)
        table.cell(i,0).paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph('')
    p=doc.add_paragraph('Observer. Vérifier. Démontrer. Documenter.', style='ADC Quote')
    p.paragraph_format.space_before=Pt(28)
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading('Table des matières',1)
    p=doc.add_paragraph()
    add_field(p,'TOC \\o "1-3" \\h \\z \\u','Cliquez avec le bouton droit, puis « Mettre à jour le champ ».')
    doc.add_page_break()


def callout(doc, title, text, kind='info'):
    color={'info':BLUE,'success':GREEN,'warning':ORANGE,'critical':RED}[kind]
    table=doc.add_table(1,2)
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width=Cm(.25); table.columns[1].width=Cm(15.8)
    set_cell_shading(table.cell(0,0),color); set_cell_shading(table.cell(0,1),'F8FAFC')
    table.cell(0,0).text=''
    p=table.cell(0,1).paragraphs[0]
    r=p.add_run(title+'\n'); r.bold=True; r.font.color.rgb=RGBColor.from_string(color)
    r=p.add_run(text); r.font.size=Pt(9.5)
    for c in table.rows[0].cells:
        set_cell_border(c,top={'val':'single','sz':'4','color':'E5E7EB'},bottom={'val':'single','sz':'4','color':'E5E7EB'})
    doc.add_paragraph('').paragraph_format.space_after=Pt(0)


def add_version_table(doc):
    doc.add_heading('Historique des versions',1)
    data=[
        ('0.1','2026-07-27','Équipe ADC Studio','Création initiale'),
        ('1.0','AAAA-MM-JJ','Auteur','Version approuvée'),
    ]
    table=doc.add_table(1,4)
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    headers=['Version','Date','Auteur','Modification']
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; c.text=h; set_cell_shading(c,BLUE)
        c.paragraphs[0].runs[0].font.bold=True; c.paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for row in data:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    for row in table.rows:
        for c in row.cells:
            set_cell_border(c,top={'val':'single','sz':'4','color':BORDER},bottom={'val':'single','sz':'4','color':BORDER},left={'val':'single','sz':'4','color':BORDER},right={'val':'single','sz':'4','color':BORDER})
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.5)
    doc.add_paragraph('')


def build_doc(path, sample=True):
    doc=Document()
    configure_styles(doc)
    for s in doc.sections: configure_page(s); add_header_footer(s)
    set_update_fields(doc)
    add_cover(doc, 'Rapport technique', 'Modèle professionnel ADC Studio', 'Client exemple', '0.2')
    add_toc(doc)
    add_version_table(doc)
    doc.add_heading('Résumé exécutif',1)
    p=doc.add_paragraph('Ce document démontre le framework Word A.D.C. : hiérarchie visuelle, styles réutilisables, composants de synthèse, tableaux et blocs techniques.',style='ADC Lead')
    callout(doc,'Conclusion principale','Le modèle fournit une base homogène pour les audits, migrations, architectures et rapports d’incident.','success')
    doc.add_heading('1. Contexte et objectif',1)
    doc.add_paragraph('Décrire ici la situation initiale, le périmètre et le résultat attendu. Les affirmations doivent être classées implicitement comme faits observés, mesures, déductions, hypothèses ou décisions.')
    doc.add_heading('1.1 Périmètre',2)
    for item in ['Infrastructure concernée','Services inclus','Exclusions explicites']:
        doc.add_paragraph(item,style='List Bullet')
    callout(doc,'Point d’attention','Une exclusion importante doit être visible avant les conclusions.','warning')
    doc.add_heading('2. Observations',1)
    doc.add_paragraph('Présenter les éléments reproductibles et vérifiables avant toute interprétation.')
    table=doc.add_table(1,4); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(['ID','Observation','Preuve','Statut']):
        table.cell(0,i).text=h; set_cell_shading(table.cell(0,i),BLUE)
        table.cell(0,i).paragraphs[0].runs[0].font.bold=True; table.cell(0,i).paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for row in [('OBS-001','Service indisponible','Journal système','Confirmé'),('OBS-002','Charge CPU stable','Mesure iLO','Confirmé')]:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    for row in table.rows:
        for c in row.cells:
            set_cell_border(c,top={'val':'single','sz':'4','color':BORDER},bottom={'val':'single','sz':'4','color':BORDER},left={'val':'single','sz':'4','color':BORDER},right={'val':'single','sz':'4','color':BORDER})
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.5)
    doc.add_paragraph('Tableau 1 - Exemple de registre d’observations',style='ADC Caption')
    doc.add_heading('3. Analyse',1)
    doc.add_heading('3.1 Hypothèse testée',2)
    doc.add_paragraph('Formuler l’hypothèse, la méthode de vérification et le résultat. Éviter de présenter une déduction comme un fait.')
    callout(doc,'Information','Les blocs sont conçus pour porter une information courte et actionnable.','info')
    doc.add_heading('3.2 Extrait technique',2)
    p=doc.add_paragraph(style='ADC Code')
    p.add_run('DBCC CHECKDB ([APPPROD_DB]) WITH TABLOCK;\nGO\n')
    set_cell=None
    doc.add_heading('4. Recommandations',1)
    recs=[('R1','Prioritaire','Tester sur un second serveur','Réduit l’incertitude sur la cause'),('R2','Planifiée','Documenter la configuration','Améliore la reproductibilité')]
    table=doc.add_table(1,4)
    for i,h in enumerate(['Réf.','Priorité','Action','Justification']):
        table.cell(0,i).text=h; set_cell_shading(table.cell(0,i),DARK)
        table.cell(0,i).paragraphs[0].runs[0].font.bold=True; table.cell(0,i).paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(WHITE)
    for row in recs:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    for row in table.rows:
        for c in row.cells:
            set_cell_border(c,top={'val':'single','sz':'4','color':BORDER},bottom={'val':'single','sz':'4','color':BORDER},left={'val':'single','sz':'4','color':BORDER},right={'val':'single','sz':'4','color':BORDER})
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.5)
    callout(doc,'Risque critique','Utiliser ce bloc uniquement lorsqu’une action ou une décision est réellement urgente.','critical')
    doc.add_heading('Annexe A - Règles d’utilisation',1)
    for item in ['Mettre à jour la table des matières avant publication.','Remplacer toutes les métadonnées de couverture.','Supprimer les exemples et textes d’aide.','Exporter en PDF après contrôle visuel.']:
        doc.add_paragraph(item,style='List Number')
    doc.core_properties.title='ADC Studio - Framework Word'
    doc.core_properties.subject='Modèle professionnel de rapport A.D.C.'
    doc.core_properties.author='A.D.C. srl - ADC Studio'
    doc.core_properties.keywords='ADC Studio, documentation, rapport, modèle Word'
    doc.save(path)


def docx_to_dotx(src, dst):
    with ZipFile(src,'r') as zin, ZipFile(dst,'w',ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data=zin.read(item.filename)
            if item.filename=='[Content_Types].xml':
                text=data.decode('utf-8')
                text=text.replace('application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml',
                                  'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml')
                data=text.encode('utf-8')
            zout.writestr(item,data)

sample=ROOT/'examples'/'ADC-Report-Example-v0.2.docx'
build_doc(sample)
template_docx=ROOT/'templates'/'word'/'ADC-Report-Template-v0.2.docx'
shutil.copy2(sample,template_docx)
docx_to_dotx(template_docx, ROOT/'templates'/'word'/'ADC-Report-Template-v0.2.dotx')

(ROOT/'docs'/'WORD_FRAMEWORK.md').write_text('''# ADC Word Framework v0.2\n\n## Objet\n\nLe framework Word fournit la base commune des rapports A.D.C. : audits, incidents, migrations, architectures et analyses techniques.\n\n## Livrables\n\n- `ADC-Report-Template-v0.2.dotx` : modèle Word installable.\n- `ADC-Report-Template-v0.2.docx` : version éditable et portable.\n- `ADC-Report-Example-v0.2.docx` : document de démonstration.\n\n## Styles disponibles\n\n- Normal et Titre 1 à Titre 4\n- ADC Subtitle\n- ADC Lead\n- ADC Metadata\n- ADC Caption\n- ADC Code\n- ADC Callout Info / Success / Warning / Critical\n- ADC Quote\n\n## Utilisation\n\n1. Copier le `.dotx` dans le dossier des modèles Office ou l’ouvrir directement.\n2. Créer un nouveau document depuis ce modèle.\n3. Remplacer les métadonnées de couverture.\n4. Utiliser exclusivement les styles fournis.\n5. Mettre à jour la table des matières et les champs avant publication.\n6. Contrôler visuellement le DOCX puis le PDF exporté.\n\n## Règle de source unique\n\nUne évolution graphique doit être apportée au modèle, jamais reproduite manuellement dans chaque rapport.\n''',encoding='utf-8')

(ROOT/'docs'/'ADR-0002-word-framework.md').write_text('''# ADR-0002 - Word Framework\n\nStatus: Accepted  \nDate: 2026-07-27\n\n## Contexte\n\nLes rapports A.D.C. doivent présenter une identité homogène sans dépendre du copier-coller d’anciens documents.\n\n## Décision\n\nAdopter un modèle Word versionné comme source unique des styles et composants de rapport. Le format `.dotx` est le livrable principal ; une version `.docx` reste fournie pour inspection et compatibilité.\n\n## Principes\n\n- styles nommés plutôt que mise en forme locale ;\n- couverture, en-têtes, pieds de page et champs communs ;\n- composants sobres et accessibles ;\n- aucune donnée client dans le modèle ;\n- exemples séparés du modèle de production ;\n- contrôle visuel obligatoire avant publication.\n\n## Conséquences\n\nLes changements d’identité visuelle sont centralisés. Les anciens rapports ne sont pas modifiés automatiquement. Toute adaptation client spécifique doit rester une extension explicite et documentée.\n''',encoding='utf-8')

(ROOT/'brand'/'word_style_specification.md').write_text('''# Spécification visuelle Word v0.2\n\n## Typographie\n\n- Titres : Aptos Display\n- Corps : Aptos\n- Code : Consolas\n\n## Couleurs\n\n- Bleu A.D.C. : `#0054A6`\n- Texte principal : `#1F2937`\n- Gris secondaire : `#6B7280`\n- Vert validation : `#2E7D32`\n- Orange attention : `#F59E0B`\n- Rouge critique : `#C62828`\n\n## Mise en page\n\n- A4 portrait\n- Marges gauche/droite : 2,1 cm\n- Marges haute/basse : environ 1,8 cm\n- Pied de page : confidentialité et pagination\n\n## Hiérarchie\n\nLa couleur sert à structurer, jamais à décorer. Le bleu identifie la marque et les titres principaux ; le rouge est réservé aux situations critiques.\n''',encoding='utf-8')

(ROOT/'tools'/'build_word_framework.py').write_text(Path('/mnt/data/build_adc_sprint002.py').read_text(encoding='utf-8'),encoding='utf-8')

(ROOT/'CHANGELOG-Sprint002.md').write_text('''# Changelog - Sprint 002\n\n## [0.2.0] - 2026-07-27\n\n### Added\n- Framework Word A.D.C. versionné.\n- Modèle `.dotx` et variante `.docx`.\n- Couverture, historique des versions, table des matières et pagination.\n- Styles nommés pour titres, corps, code, citations et métadonnées.\n- Composants d’information, validation, attention et criticité.\n- Exemple complet de rapport.\n- ADR-0002 et spécification visuelle Word.\n- Script reproductible de génération.\n''',encoding='utf-8')
(ROOT/'VERSION').write_text('0.2.0\n',encoding='utf-8')
(ROOT/'COMMIT_MESSAGE.txt').write_text('Add ADC Word framework and reusable report components\n',encoding='utf-8')

zip_path=Path('/mnt/data/ADC-Studio-Sprint002.zip')
with ZipFile(zip_path,'w',ZIP_DEFLATED) as z:
    for p in ROOT.rglob('*'):
        if p.is_file(): z.write(p,p.relative_to(ROOT.parent))
print(zip_path)
