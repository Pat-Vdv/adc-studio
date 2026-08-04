"""Tests du renderer DOCX (increment 2).

Périmètre : rendu de la Cover (C-001-cover), de l'identité documentaire
(C-002-identity-page), du résumé exécutif (C-003-executive-summary) et de
l'environnement (C-009-environment) depuis le Document IR.
"""
from __future__ import annotations

import copy
import json
from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from adc_engine import ComponentInstance, compose_document
from adc_engine.render_docx import render_docx

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "reference_reports" / "incident_report" / "data" / "sql_server_2014_incident.json"


def _data() -> dict:
    return json.loads(DATA.read_text(encoding="utf-8-sig"))


def _texts(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def test_render_creates_file(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    assert out.exists()
    assert out.stat().st_size > 0


def test_cover_texts_present(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Investigation — Blocage SQL Server lors de DBCC CHECKDB" in text
    assert "Soc01" in text
    assert "Confidentiel" in text
    assert "0.1-draft" in text  # version
    assert "2026-07-28" in text  # date
    assert "A.D.C. srl" in text  # auteur
    assert "ADC-SOC01-2026-SQL2014-001" in text  # référence


def test_identity_page_texts_present(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Identité du document" in text
    assert "fr-BE" in text  # langue, propre à la page d'identité


_REVISIONS_HEADER = ["Version", "Date", "Auteur", "Objet"]


def _table_rows(path: Path, header: list[str]) -> list[list[str]]:
    """Lignes de la table portant cet en-tête (liste vide si aucune)."""
    for table in DocxDocument(str(path)).tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows and rows[0] == header:
            return rows
    return []


def test_identity_page_omits_empty_optional_tables(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Révisions" not in text
    assert "Diffusion" not in text
    assert _table_rows(out, _REVISIONS_HEADER) == []


def test_identity_page_renders_optional_tables(tmp_path):
    data = _data()
    data["report"]["revisions"] = [
        {"version": "0.1-draft", "date": "2026-07-28", "author": "A.D.C. srl", "summary": "Création"}
    ]
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    assert "Révisions" in "\n".join(p.text for p in DocxDocument(str(out)).paragraphs)
    rows = _table_rows(out, _REVISIONS_HEADER)
    assert rows[1:] == [["0.1-draft", "2026-07-28", "A.D.C. srl", "Création"]]


def test_executive_summary_renders_payload_faithfully(tmp_path):
    # Le rendu est comparé au payload composé, pas au texte de la source : la
    # rédaction du rapport de référence n'a aucune incidence sur ce test.
    data = _data()
    data["executive_summary"]["context"] = "Contexte du blocage.\n\nSeconde partie du contexte."
    document = compose_document(data)
    summary = document.components[2]
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]

    assert summary.payload["heading"] in paragraphs
    for key, heading in (
        ("context", "Contexte"),
        ("business_impact", "Impact métier"),
        ("conclusion", "Conclusion"),
        ("recommended_action", "Action recommandée"),
    ):
        blocks = summary.payload[key]
        assert (heading in paragraphs) is bool(blocks)  # titre ssi volet renseigné
        for block in blocks:
            assert block in paragraphs


def test_executive_summary_omits_empty_sections(tmp_path):
    data = _data()
    data["executive_summary"] = {"context": "Seul volet renseigné."}
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Contexte" in paragraphs
    assert "Impact métier" not in paragraphs  # pas de titre sans contenu
    assert "Action recommandée" not in paragraphs


def test_environment_texts_present(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Environnement" in text
    assert "SRV-SQL-01" in text
    assert "Microsoft SQL Server 2014 Standard" in text
    assert "French_CI_AS" in text
    assert "40" in text  # processeurs logiques, valeur numérique rendue


def test_environment_storage_table(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    assert "Stockage" in _texts(out)
    rows = _table_rows(out, ["Volume", "Rôle", "Unité d'allocation (Ko)"])
    assert rows[1] == ["C:", "Système", ""]  # unité non renseignée : cellule vide
    assert rows[2] == ["D:", "SQL", "64"]
    assert [row[0] for row in rows[1:]] == ["C:", "D:", "E:", "I:"]


def test_environment_without_storage_has_no_table(tmp_path):
    data = _data()
    data["environment"].pop("storage")
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    assert "Stockage" not in _texts(out)
    assert _table_rows(out, ["Volume", "Rôle", "Unité d'allocation (Ko)"]) == []


_TIMELINE_HEADER = ["Horodatage", "Événement", "Description"]


def test_timeline_table_follows_source_order(tmp_path):
    data = _data()
    data["timeline"].append(
        {
            "id": "timeline-000",
            "timestamp": "2026-07-22",
            "title": "Signalement",
            "description": "Appel du client.",
        }
    )
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    assert "Chronologie" in _texts(out)
    rows = _table_rows(out, _TIMELINE_HEADER)
    assert rows[1] == [
        "2026-07-23",
        "Sauvegardes préalables",
        "Sauvegardes complètes et VERIFYONLY réalisés avant intervention.",
    ]
    assert [row[0] for row in rows[1:]] == ["2026-07-23", "2026-07-22"]  # ordre source


def test_timeline_empty_renders_no_section(tmp_path):
    # Instance présente mais sans entrée : ni titre ni tableau vide.
    data = _data()
    data["timeline"] = ["entrée invalide ignorée"]
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    assert "Chronologie" not in _texts(out)
    assert _table_rows(out, _TIMELINE_HEADER) == []


def test_timeline_absent_renders_no_section(tmp_path):
    data = _data()
    data.pop("timeline")
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    assert "Chronologie" not in _texts(out)
    assert "Environnement" in _texts(out)  # le reste du rapport est intact


def test_incident_context_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["incident_context"]["description"] = "Circonstances.\n\nSeconde partie."
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Contexte de l'incident" in paragraphs
    assert "Déclencheur : Exécution de DBCC CHECKDB" in paragraphs
    assert "Périmètre : Instance SQL Server APPPROD" in paragraphs
    assert "Circonstances." in paragraphs
    assert "Seconde partie." in paragraphs


def test_incident_context_status_is_french_in_the_docx_only(tmp_path):
    document = compose_document(_data())
    context = next(c for c in document.components if c.instance_id == "incident-context")
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert context.payload["status"] == "investigated"  # IR canonique
    assert "Statut : Investigé" in text
    assert "investigated" not in text


def test_incident_context_unknown_status_is_rendered_verbatim(tmp_path):
    data = _data()
    data["incident_context"]["status"] = "reopened"
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    assert "Statut : reopened" in _texts(out)


def test_incident_context_omits_absent_qualifiers(tmp_path):
    data = _data()
    data["incident_context"] = {"description": "Seules les circonstances."}
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    # Bornée à la section : les blocs suivants portent aussi un statut.
    start = paragraphs.index("Contexte de l'incident")
    section = paragraphs[start : paragraphs.index("Environnement", start)]
    assert "Seules les circonstances." in section
    assert not any(p.startswith(("Déclencheur", "Périmètre", "Statut")) for p in section)


def test_incident_context_absent_renders_no_section(tmp_path):
    data = _data()
    data.pop("incident_context")
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    text = _texts(out)
    assert "Contexte de l'incident" not in text
    assert "Environnement" in text  # le reste du rapport est intact


def test_investigation_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["investigations"][0]["description"] = "Contrôles exécutés.\n\nSeconde partie."
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Investigation — Analyse de l’environnement SQL Server" in paragraphs
    assert "Résultat : TODO" in paragraphs  # valeur déclarée, non traduite
    assert "Contrôles exécutés." in paragraphs
    assert "Seconde partie." in paragraphs
    assert "Investigations" not in paragraphs  # aucun titre de partie commun


def test_investigations_are_rendered_in_source_order(tmp_path):
    data = _data()
    data["investigations"].append(
        {
            "id": "investigation-000",
            "title": "A — Analyse préliminaire",
            "description": "Contrôles préalables.",
            "result": "inconclusive",
        }
    )
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    first = paragraphs.index("Investigation — Analyse de l’environnement SQL Server")
    second = paragraphs.index("Investigation — A — Analyse préliminaire")
    assert first < second  # ni identifiant ni titre ne réordonnent le rendu


def test_investigation_omits_absent_result(tmp_path):
    data = _data()
    data["investigations"] = [{"id": "investigation-001", "title": "Sans détail"}]
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    start = paragraphs.index("Investigation — Sans détail")
    section = paragraphs[start : paragraphs.index("Constat — Blocage observé pendant DBCC CHECKDB")]
    assert not any(p.startswith("Résultat") for p in section)


def test_investigation_does_not_leak_identifiers(tmp_path):
    out = render_docx(compose_document(_data()), tmp_path / "report.docx")
    assert "investigation-001" not in _texts(out)


def test_investigations_absent_render_no_section(tmp_path):
    data = _data()
    data.pop("investigations")
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    # Le titre du rapport commence lui aussi par « Investigation — » : on cible
    # donc la section, pas une sous-chaîne du document.
    assert "Investigation — Analyse de l’environnement SQL Server" not in paragraphs
    assert "Constat — Blocage observé pendant DBCC CHECKDB" in paragraphs  # reste intact


def test_probable_cause_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["probable_cause"]["statement"] = "Hypothèse principale.\n\nSeconde partie."
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Cause probable" in paragraphs
    assert "Hypothèse principale." in paragraphs
    assert "Seconde partie." in paragraphs
    assert "Constats à l'appui : Blocage observé pendant DBCC CHECKDB" in paragraphs


def test_probable_cause_confidence_is_french_in_the_docx_only(tmp_path):
    document = compose_document(_data())
    cause = next(c for c in document.components if c.instance_id == "probable-cause")
    text = _texts(render_docx(document, tmp_path / "report.docx"))
    assert cause.payload["confidence"] == "unknown"  # IR canonique
    # L'absence de certitude est une information : la ligne reste affichée.
    assert "Confiance : Indéterminée" in text
    assert "unknown" not in text


def test_probable_cause_unmapped_confidence_is_rendered_verbatim(tmp_path):
    data = _data()
    data["probable_cause"]["confidence"] = "corroborated"
    text = _texts(render_docx(compose_document(data), tmp_path / "report.docx"))
    assert "Confiance : corroborated" in text


def test_probable_cause_renders_titles_in_source_order_with_duplicates(tmp_path):
    data = _data()
    data["findings"].append(
        {
            "id": "finding-002",
            "title": "Second constat",
            "severity": "low",
            "observation": "…",
            "impact": "…",
            "analysis": "…",
            "evidence_ids": [],
        }
    )
    data["probable_cause"]["supporting_finding_ids"] = ["finding-002", "finding-001", "finding-002"]
    text = _texts(render_docx(compose_document(data), tmp_path / "report.docx"))
    assert (
        "Constats à l'appui : Second constat ; Blocage observé pendant DBCC CHECKDB ; Second constat"
        in text
    )


def test_probable_cause_unknown_reference_is_not_displayed(tmp_path):
    data = _data()
    data["probable_cause"]["supporting_finding_ids"] = ["finding-404"]
    text = _texts(render_docx(compose_document(data), tmp_path / "report.docx"))
    assert "Cause probable" in text
    assert "finding-404" not in text
    assert "Constats à l'appui" not in text


def test_probable_cause_absent_renders_no_section(tmp_path):
    data = _data()
    data.pop("probable_cause")
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Cause probable" not in paragraphs
    assert "Constat — Blocage observé pendant DBCC CHECKDB" in paragraphs  # reste intact


def _headings(path: Path, level: str) -> list[str]:
    """Titres du document au niveau demandé, dans l'ordre du document."""
    return [p.text for p in DocxDocument(str(path)).paragraphs if p.style.name == level]


def test_conclusion_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["conclusion"] = "Conclusion finale.\n\nSeconde partie."
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    # Titre de niveau 1, distinct du volet « Conclusion » du résumé exécutif,
    # qui est un sous-titre de niveau 2.
    assert "Conclusion" in _headings(out, "Heading 1")
    assert "Conclusion" in _headings(out, "Heading 2")
    assert "Conclusion finale." in paragraphs
    assert "Seconde partie." in paragraphs


def test_conclusion_absent_renders_no_section(tmp_path):
    data = _data()
    data.pop("conclusion")
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    assert "Conclusion" not in _headings(out, "Heading 1")
    assert "Preuve — État et configuration de l’environnement SQL" in [
        p.text for p in DocxDocument(str(out)).paragraphs
    ]  # le reste du rapport est intact


def test_finding_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["findings"][0]["observation"] = "Symptôme attesté.\n\nSecond paragraphe."
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Constat — Blocage observé pendant DBCC CHECKDB" in paragraphs
    assert "Observation" in paragraphs
    assert "Symptôme attesté." in paragraphs
    assert "Second paragraphe." in paragraphs
    assert "Analyse" in paragraphs


def test_finding_renders_evidence_titles_never_ids(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Preuves : État et configuration de l’environnement SQL" in text
    assert "evidence-001" not in text  # identifiant technique jamais affiché
    assert "finding-001" not in text


def test_finding_omits_evidence_line_without_references(tmp_path):
    data = _data()
    data["findings"][0]["evidence_ids"] = []
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Constat — Blocage observé pendant DBCC CHECKDB" in text
    assert "Preuves" not in text


def test_findings_are_rendered_independently(tmp_path):
    data = _data()
    data["findings"].append(
        {
            "id": "finding-002",
            "title": "Second constat",
            "severity": "low",
            "observation": "Observation du second constat.",
            "impact": "Impact mineur.",
            "analysis": "Analyse du second constat.",
            "evidence_ids": ["evidence-001"],
        }
    )
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    # Deux sections autonomes, aucun titre de partie « Constats » commun.
    assert "Constat — Blocage observé pendant DBCC CHECKDB" in paragraphs
    assert "Constat — Second constat" in paragraphs
    assert "Constats" not in paragraphs


def test_decision_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["actions_taken"][0]["description"] = "Première partie.\n\nSeconde partie."
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Mesure prise — Sécurisation préalable par sauvegarde" in paragraphs
    assert "Statut : completed" in paragraphs  # valeur canonique, non traduite
    assert "Première partie." in paragraphs
    assert "Seconde partie." in paragraphs


def test_decision_omits_absent_blocks(tmp_path):
    data = _data()
    data["actions_taken"] = [{"id": "decision-001", "title": "Mesure sans détail"}]
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Mesure prise — Mesure sans détail" in paragraphs
    # Assertion locale à la section : d'autres blocs portent aussi un statut.
    section = paragraphs[paragraphs.index("Mesure prise — Mesure sans détail") :]
    assert not any(p.startswith("Statut") for p in section)


def test_decisions_are_rendered_independently(tmp_path):
    data = _data()
    data["actions_taken"].append(
        {
            "id": "decision-002",
            "title": "Seconde mesure",
            "description": "Description de la seconde mesure.",
            "status": "in_progress",
        }
    )
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Mesure prise — Sécurisation préalable par sauvegarde" in paragraphs
    assert "Mesure prise — Seconde mesure" in paragraphs
    assert "Mesures prises" not in paragraphs  # aucun titre de partie commun


def test_recommendation_is_rendered_as_its_own_section(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Recommandation — Exécuter DBCC CHECKDB hors production" in paragraphs
    assert "Description" in paragraphs
    assert "Justification" in paragraphs
    assert "Recommandations" not in paragraphs  # aucun titre de partie commun


def test_recommendation_renders_finding_titles_never_ids(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Constats liés : Blocage observé pendant DBCC CHECKDB" in text
    assert "finding-001" not in text
    assert "recommendation-001" not in text


def test_enumerations_are_english_in_the_ir_and_french_in_the_docx(tmp_path):
    # Règle figée : valeurs canoniques anglaises dans le Document (IR),
    # libellés français produits uniquement par la couche DOCX.
    data = _data()
    data["findings"][0]["severity"] = "critical"
    data["recommendations"][0]["priority"] = "medium"
    document = compose_document(data)

    payloads = [c.payload for c in document.components]
    assert any(p.get("severity") == "critical" for p in payloads)
    assert any(p.get("priority") == "medium" for p in payloads)
    assert not any("Critique" in str(p) or "Moyenne" in str(p) for p in payloads)

    text = _texts(render_docx(document, tmp_path / "report.docx"))
    assert "Gravité : Critique" in text
    assert "Priorité : Moyenne" in text
    assert "critical" not in text
    assert "medium" not in text


def test_enumeration_labels_agree_with_their_field(tmp_path):
    # Même valeur canonique, accord différent selon le substantif affiché.
    data = _data()
    data["findings"][0]["severity"] = "high"
    data["recommendations"][0]["priority"] = "medium"
    data["risks"][0]["level"] = "high"
    text = _texts(render_docx(compose_document(data), tmp_path / "report.docx"))
    assert "Gravité : Élevée" in text
    assert "Priorité : Moyenne" in text
    assert "Niveau : Élevé" in text
    assert "Niveau : Élevée" not in text


def test_masculine_vocabulary_covers_every_level(tmp_path):
    expected = {"low": "Faible", "medium": "Moyen", "high": "Élevé", "critical": "Critique"}
    for value, label in expected.items():
        data = _data()
        data["risks"][0]["level"] = value
        text = _texts(render_docx(compose_document(data), tmp_path / f"report-{value}.docx"))
        assert f"Niveau : {label}" in text


def test_unknown_enumeration_value_is_rendered_verbatim(tmp_path):
    # Hors vocabulaire : on restitue la valeur de la source plutôt que
    # d'inventer une traduction.
    data = _data()
    data["recommendations"][0]["priority"] = "blocking"
    document = compose_document(data)
    text = _texts(render_docx(document, tmp_path / "report.docx"))
    assert "Priorité : blocking" in text


def test_recommendations_are_rendered_in_source_order(tmp_path):
    data = _data()
    data["recommendations"].append(
        {
            "id": "recommendation-002",
            "title": "Priorité critique déclarée en second",
            "priority": "critical",
            "description": "Description.",
            "rationale": "Justification.",
            "related_finding_ids": [],
        }
    )
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    first = paragraphs.index("Recommandation — Exécuter DBCC CHECKDB hors production")
    second = paragraphs.index("Recommandation — Priorité critique déclarée en second")
    assert first < second  # aucun regroupement ni tri par priorité


def test_risk_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["risks"][0]["description"] = "Risque résiduel.\n\nSeconde partie."
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Risque — Nouvelle indisponibilité lors d’un contrôle lourd en production" in paragraphs
    assert "Niveau : Élevé" in paragraphs  # accordé au substantif, IR en anglais
    assert "Risque résiduel." in paragraphs
    assert "Seconde partie." in paragraphs
    assert "Risques" not in paragraphs  # aucun titre de partie commun


def test_risk_renders_recommendation_titles_never_ids(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Traitement prévu : Exécuter DBCC CHECKDB hors production" in text
    assert "recommendation-001" not in text
    assert "risk-001" not in text


def test_risk_unknown_reference_is_not_displayed(tmp_path):
    # Référence sans libellé : ni exception, ni identifiant technique dans le
    # document, ni libellé inventé — elle sort en diagnostic côté composition.
    data = _data()
    data["risks"][0]["mitigation_recommendation_ids"] = ["recommendation-404"]
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Risque — Nouvelle indisponibilité lors d’un contrôle lourd en production" in text
    assert "recommendation-404" not in text
    assert "Traitement prévu" not in text


def test_risks_are_rendered_in_source_order(tmp_path):
    data = _data()
    data["risks"].append(
        {
            "id": "risk-002",
            "title": "Niveau critique déclaré en second",
            "level": "critical",
            "description": "Description.",
            "mitigation_recommendation_ids": [],
        }
    )
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    first = paragraphs.index(
        "Risque — Nouvelle indisponibilité lors d’un contrôle lourd en production"
    )
    second = paragraphs.index("Risque — Niveau critique déclaré en second")
    assert first < second


def test_evidence_is_rendered_as_its_own_section(tmp_path):
    data = _data()
    data["evidence"][0]["content"] = "Première partie.\n\nSeconde partie."
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Preuve — État et configuration de l’environnement SQL" in paragraphs
    assert "Nature : technical_observation" in paragraphs  # valeur déclarée
    assert "Origine : Serveur SRV-SQL-01" in paragraphs
    assert "Contenu" in paragraphs
    assert "Première partie." in paragraphs
    assert "Seconde partie." in paragraphs
    assert "Preuves" not in paragraphs  # aucun titre de partie commun


def test_evidence_omits_absent_blocks(tmp_path):
    data = _data()
    data["evidence"] = [{"id": "evidence-001", "title": "Preuve sans détail"}]
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "Preuve — Preuve sans détail" in paragraphs
    assert not any(p.startswith("Nature") or p.startswith("Origine") for p in paragraphs)
    assert "Contenu" not in paragraphs


def test_evidence_is_rendered_in_source_order(tmp_path):
    data = _data()
    data["evidence"].append(
        {
            "id": "evidence-002",
            "title": "Seconde preuve",
            "kind": "log_extract",
            "description": "Description.",
            "content": "Extrait.",
            "source": "Instance APPPROD",
        }
    )
    document = compose_document(data)
    out = render_docx(document, tmp_path / "report.docx")
    paragraphs = [p.text for p in DocxDocument(str(out)).paragraphs]
    first = paragraphs.index("Preuve — État et configuration de l’environnement SQL")
    second = paragraphs.index("Preuve — Seconde preuve")
    assert first < second


def test_evidence_section_does_not_leak_identifiers(tmp_path):
    document = compose_document(_data())
    out = render_docx(document, tmp_path / "report.docx")
    assert "evidence-001" not in _texts(out)


def test_unsupported_component_is_skipped_not_crashed(tmp_path):
    document = compose_document(_data())
    # Injecte une instance sans renderer : le rendu ne doit ni planter ni la rendre.
    document = replace(
        document,
        components=document.components
        + (ComponentInstance("C-999-unknown", "ghost", {"title": "NE DOIT PAS APPARAITRE"}),),
    )
    out = render_docx(document, tmp_path / "report.docx")
    text = _texts(out)
    assert "Soc01" in text  # la Cover est bien rendue
    assert "NE DOIT PAS APPARAITRE" not in text  # le composant inconnu est ignoré


def test_render_does_not_mutate_document(tmp_path):
    document = compose_document(_data())
    before = copy.deepcopy(document)
    render_docx(document, tmp_path / "report.docx")
    assert document == before


# --- Blocs narratifs vides : une règle, pas quatre -------------------------
#
# Les quatre blocs suivaient deux conventions opposées : la conclusion se
# gardait du titre orphelin, le contexte et la cause probable le posaient quand
# même. Le lecteur voyait donc un intitulé sans contenu dans un cas, et rien
# dans l'autre, pour un même état de la source.
#
# La convention retenue est celle que le renderer suivait déjà pour les volets
# d'un constat, d'une preuve et de la conclusion : un bloc sans contenu n'écrit
# rien. Elle relève du rendu (ADR-0009, I3) et d'aucun contrat de source.

_EMPTY = (
    ("incident_context", {}, "Contexte de l'incident"),
    ("probable_cause", {}, "Cause probable"),
    ("conclusion", "", "Conclusion"),
    ("investigations", [{"id": "inv-vide"}], "Investigation"),
)

_MINIMAL = (
    ("incident_context", {"scope": "Instance de production"}, "Contexte de l'incident"),
    ("probable_cause", {"confidence": "unknown"}, "Cause probable"),
    ("conclusion", "Un mot suffit.", "Conclusion"),
    ("investigations", [{"id": "inv-1", "result": "Confirmé"}], "Investigation"),
)


@pytest.mark.parametrize("node,empty,heading", _EMPTY, ids=[case[0] for case in _EMPTY])
def test_an_empty_narrative_block_writes_no_orphan_heading(tmp_path, node, empty, heading):
    data = {**_data(), node: empty}
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    assert heading not in _headings(out, "Heading 1")


@pytest.mark.parametrize("node,content,heading", _MINIMAL, ids=[case[0] for case in _MINIMAL])
def test_a_single_field_is_enough_to_render_the_block(tmp_path, node, content, heading):
    # L'autre moitié de la règle : la garde ne doit pas manger un bloc qui a
    # quelque chose à dire, si peu que ce soit.
    data = {**_data(), node: content}
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    assert heading in _headings(out, "Heading 1")


def test_a_block_whose_only_references_are_unresolved_writes_nothing(tmp_path):
    """Le vide se juge sur ce qui serait affiché, pas sur les clés du payload.

    Une référence sans libellé n'est jamais écrite (I4) : un bloc qui n'aurait
    que celle-là n'a rien à montrer, et ne pose donc pas son titre. L'écart, lui,
    reste dit — par la composition et par le validateur métier.
    """
    data = {**_data(), "probable_cause": {"supporting_finding_ids": ["finding-404"]}}
    out = render_docx(compose_document(data), tmp_path / "report.docx")
    assert "Cause probable" not in _headings(out, "Heading 1")
    assert "finding-404" not in _texts(out)
